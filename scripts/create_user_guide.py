from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Title
title = doc.add_heading('Tài liệu Hướng dẫn Sử dụng', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

sub = doc.add_paragraph('Learnify Studio — Nền tảng tạo video bài giảng AI')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
sub.runs[0].font.size = Pt(12)

doc.add_paragraph()

# Overview
doc.add_heading('1. Tổng quan', level=1)
p = doc.add_paragraph(
    'Learnify Studio là nền tảng tạo video bài giảng AI dành cho giảng viên. '
    'Hệ thống tự động sinh script, chọn hình ảnh, và xuất video từ nội dung bài giảng của bạn.'
)

# Login
doc.add_heading('2. Đăng nhập', level=1)
doc.add_paragraph('• Truy cập Learnify Studio qua link được cấp bởi Admin')
doc.add_paragraph('• Đăng nhập bằng tài khoản Giảng viên hoặc Admin')
doc.add_paragraph('• Sau khi đăng nhập, bạn vào thẳng Tổng quan (Dashboard)')

# Dashboard
doc.add_heading('3. Dashboard — Tổng quan', level=1)

table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Thẻ thông tin'
hdr[1].text = 'Ý nghĩa'
for cell in hdr:
    cell.paragraphs[0].runs[0].font.bold = True

rows = [
    ('Tổng video', 'Số video đã tạo'),
    ('Hoàn thành', 'Video đã xuất thành công'),
    ('Đang xử lý', 'Video đang được render'),
    ('Chi phí tháng', 'Chi phí GPU phát sinh trong tháng'),
]
for i, (k, v) in enumerate(rows):
    row = table.rows[i + 1].cells
    row[0].text = k
    row[1].text = v

# Create Video
doc.add_heading('4. Tạo video mới', level=1)
doc.add_paragraph('Vào Tạo video mới trên thanh điều hướng. Quy trình gồm 4 bước:')

doc.add_heading('Bước 1 — Nhập thông tin', level=2)
doc.add_paragraph('Có 2 cách nhập nội dung:')

doc.add_paragraph('✏️ Nhập tay:', style='List Bullet')
items = ['Môn học / Chủ đề', 'Outline bài giảng (vài dòng, AI tự mở rộng)', 'Thời lượng: 5 / 10 / 15 / 20 / 30 / 45 phút']
for item in items:
    p = doc.add_paragraph(item, style='List Bullet 2')

doc.add_paragraph('📎 Upload file:', style='List Bullet')
for item in ['Upload PDF, PPTX, DOCX — AI đọc và sinh script từ nội dung', 'Tối đa 20MB']:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_heading('Bước 2 — Sinh Script (AI)', level=2)
for item in [
    'Click Sinh Script → AI tự tạo script chi tiết từng slide',
    'Mỗi slide có: nội dung giọng đọc (GV đọc) và nội dung slide (hiển thị)',
    'AI tự gợi ý hình ảnh phù hợp từng slide',
    'Có thể chỉnh sửa trực tiếp từng slide, thêm/xóa slide',
    'Upload ảnh tùy chỉnh cho từng slide nếu muốn',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Bước 3 — Chọn kiểu video', level=2)

table2 = doc.add_table(rows=3, cols=2)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Kiểu video'
hdr2[1].text = 'Mô tả'
for cell in hdr2:
    cell.paragraphs[0].runs[0].font.bold = True

tiers = [
    ('🔴 Tier 1 — Avatar AI', 'Mặt GV thật nói trong video. Phù hợp intro, video marketing.'),
    ('🟡 Tier 2 — Slide + Voice (Khuyến nghị)', 'Slide đẹp + giọng đọc AI. Tối ưu cho nội dung lý thuyết chính.'),
]
for i, (k, v) in enumerate(tiers):
    row = table2.rows[i + 1].cells
    row[0].text = k
    row[1].text = v

doc.add_paragraph()
doc.add_paragraph('Tier 1 yêu cầu upload ảnh giảng viên và nhập đoạn intro.')

doc.add_paragraph('Chọn giọng đọc:')
for v in ['Edge TTS: Hoài My (nữ), Nam Minh (nam)', 'ElevenLabs: Thảo, Ninh Đôn, Hiện, Thắm, Nhật', 'Google TTS (miễn phí)']:
    doc.add_paragraph(v, style='List Bullet')

doc.add_paragraph('Chọn theme màu slide: Dark, Ocean, Midnight, Forest, Sunset, Light')

doc.add_heading('Bước 4 — Xác nhận & Submit', level=2)
doc.add_paragraph('Xem lại toàn bộ thông tin rồi click Tạo video. Job sẽ được gửi lên hàng đợi.')

# Track videos
doc.add_heading('5. Theo dõi Video', level=1)
doc.add_paragraph('Vào Video bài giảng để xem danh sách:')

table3 = doc.add_table(rows=4, cols=2)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Trạng thái'
hdr3[1].text = 'Ý nghĩa'
for cell in hdr3:
    cell.paragraphs[0].runs[0].font.bold = True

statuses = [
    ('🔵 Đang xử lý', 'Video đang render'),
    ('🟢 Hoàn thành', 'Sẵn sàng xem / tải'),
    ('🔴 Lỗi', 'Render thất bại — liên hệ Admin'),
]
for i, (k, v) in enumerate(statuses):
    row = table3.rows[i + 1].cells
    row[0].text = k
    row[1].text = v

# Profile
doc.add_heading('6. Hồ sơ Giảng viên', level=1)
for item in ['Tên giảng viên', 'Ảnh đại diện (dùng cho Tier 1 Avatar)', 'Thông tin liên hệ']:
    doc.add_paragraph(item, style='List Bullet')

# Admin
doc.add_heading('7. Quản trị (Admin)', level=1)
doc.add_paragraph('Chỉ hiển thị với tài khoản Admin:')
for item in ['Xem tất cả giảng viên trong tổ chức', 'Xem tổng video và chi phí toàn bộ hệ thống', 'Quản lý tài khoản']:
    doc.add_paragraph(item, style='List Bullet')

# Note
doc.add_heading('Lưu ý', level=1)
for note in [
    'Tier 2 thích hợp cho 80% nội dung bài giảng — chi phí thấp, tối ưu.',
    'Tier 1 dùng cho video intro, marketing khóa học — chuyên nghiệp hơn.',
    'Video được lưu trên cloud, truy cập mọi lúc qua dashboard.',
]:
    doc.add_paragraph(note, style='List Bullet')

out_path = r"C:\Users\admin\.gemini\antigravity\brain\8571f8ef-d595-4c87-8ac6-bf3bdf7a92a5\Learnify_Studio_UserGuide.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
